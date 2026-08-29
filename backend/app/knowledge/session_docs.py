"""会话级文档智能。

策略 A：本地抽取正文为主；仅 fileid 能力模型才注入 fileid://。
小文档 mode=fulltext（全文 system 注入 + session_chunks）；
大文档/图片 mode=session_rag（仅 session_chunks）。
扫描 PDF：MinerU 过短且 VISION_PDF_ENABLED 时回退 Vision PDF。
"""
from __future__ import annotations

import asyncio
import hashlib
from pathlib import Path
from typing import Iterable

from loguru import logger
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.db.models import SessionFile
from app.knowledge.contracts import IMAGE_EXTENSIONS, build_child_chunk_id
from app.knowledge.indexer import (
    IndexableChunk,
    ensure_collection,
    upsert_chunks,
)
from app.knowledge.ingestion.chunker import chunk_document
from app.knowledge.ingestion.parser import (
    SESSION_UPLOAD_EXTENSIONS,
    ParsedDocument,
    Section,
    parse_document,
    parse_image_document,
    parse_scanned_pdf_document,
)
from app.llm.dashscope import dashscope_client

_INLINE_TEXT_EXTS = {".md", ".markdown", ".txt", ".text"}
_FULLTEXT_MODES = frozenset({"fulltext", "files_api"})  # files_api = legacy alias
_INDEXED_MODES = frozenset({"fulltext", "files_api", "session_rag"})


def _inline_max_chars() -> int:
    return max(1_000, int(settings.session_inline_max_chars))


def _fileid_capable() -> bool:
    model = (settings.chat_model or "").strip().lower()
    if not model:
        return False
    allowed = [
        m.strip().lower()
        for m in (settings.fileid_capable_models or "").split(",")
        if m.strip()
    ]
    for a in allowed:
        if model == a or model.startswith(f"{a}-") or model.startswith(f"{a}."):
            return True
    return False


def _estimate_tokens(text: str) -> int:
    if not text:
        return 0
    chinese = sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")
    other = len(text) - chinese
    return chinese + max(1, other // 4)


def _file_sha(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def _wants_fulltext(path: Path, extracted_text: str | None, ext: str) -> bool:
    """Small enough for full system-message injection (not images)."""
    if ext in IMAGE_EXTENSIONS:
        return False
    if ext in _INLINE_TEXT_EXTS:
        text = extracted_text or path.read_text(encoding="utf-8", errors="ignore")
        return _estimate_tokens(text) <= settings.files_api_inline_max_tokens
    if ext in {".pdf", ".docx"}:
        if extracted_text is not None:
            return _estimate_tokens(extracted_text) <= settings.files_api_inline_max_tokens
        return path.stat().st_size <= settings.files_api_inline_max_bytes
    return False


def preview_kind_for_filename(filename: str) -> str | None:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in IMAGE_EXTENSIONS:
        return "image"
    if ext in _INLINE_TEXT_EXTS | {".docx"}:
        return "text"
    return None


def _parsed_from_plain_text(title: str, text: str, *, parser: str, fmt: str) -> ParsedDocument:
    body = text.strip()
    if not body:
        raise ValueError("Document contains no extractable text.")
    return ParsedDocument(
        title=title,
        sections=[Section(heading_path=title, text=body)],
        metadata={"format": fmt, "parser": parser},
    )


def _extract_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for table_row in table.rows:
            cells = [(c.text or "").strip() for c in table_row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts).strip()


def _extract_attachment_text(path: Path) -> str | None:
    """Sync extract for txt/md/docx (and MinerU PDF). Prefer for preview/backfill."""
    if not path.is_file():
        return None
    ext = path.suffix.lower()
    try:
        if ext in _INLINE_TEXT_EXTS:
            text = path.read_text(encoding="utf-8", errors="ignore").strip()
            return text or None
        if ext == ".docx":
            text = _extract_docx_text(path)
            return text or None
        if ext == ".pdf":
            text = parse_document(path).full_text().strip()
            return text or None
    except Exception as e:
        logger.warning("Failed to extract inline text from {}: {}", path.name, e)
        return None
    return None


async def _parse_session_document(
    storage_path: Path, *, ext: str, original_filename: str
) -> tuple[ParsedDocument, str]:
    """Parse session upload; PDF may fall back to Vision when MinerU text is too short."""
    if ext in IMAGE_EXTENSIONS:
        parsed = await parse_image_document(storage_path)
        return parsed, (parsed.metadata or {}).get("parser") or "vision"

    if ext == ".docx":
        # Prefer lightweight extract for speed; fall back to MinerU if empty.
        text = _extract_docx_text(storage_path)
        if text:
            return (
                _parsed_from_plain_text(storage_path.stem, text, parser="python-docx", fmt="docx"),
                "python-docx",
            )
        parsed = parse_document(storage_path)
        return parsed, (parsed.metadata or {}).get("parser") or "mineru"

    if ext in _INLINE_TEXT_EXTS:
        parsed = parse_document(storage_path)
        return parsed, (parsed.metadata or {}).get("parser") or "direct"

    if ext == ".pdf":
        parser_used = "mineru"
        try:
            parsed = parse_document(storage_path)
        except Exception as e:
            logger.warning("MinerU failed for {}: {}", original_filename, e)
            parsed = None  # type: ignore[assignment]
            if not settings.vision_pdf_enabled:
                raise

        if parsed is not None:
            full = parsed.full_text().strip()
            if len(full) >= settings.session_min_extract_chars:
                return parsed, (parsed.metadata or {}).get("parser") or "mineru"
            if not settings.vision_pdf_enabled:
                raise ValueError(
                    f"Extracted text too short ({len(full)} chars) for {original_filename}. "
                    "若为扫描件请确认 MINERU_OCR=true 或开启 VISION_PDF_ENABLED。"
                )
            logger.info(
                "Session PDF {} short extract ({} chars) → Vision PDF fallback",
                original_filename,
                len(full),
            )
        else:
            logger.info("Session PDF {} MinerU failed → Vision PDF fallback", original_filename)

        parsed = await parse_scanned_pdf_document(storage_path)
        parser_used = "vision_pdf"
        full = parsed.full_text().strip()
        if len(full) < settings.session_min_extract_chars:
            raise ValueError(
                f"Extracted text too short ({len(full)} chars) for {original_filename} "
                f"after Vision PDF. 请检查 VISION_MODEL / 页数上限。"
            )
        return parsed, parser_used

    raise ValueError(f"Unsupported session upload type: {ext}")


async def _index_parsed_to_session(
    *,
    thread_id: str,
    sha: str,
    parsed: ParsedDocument,
) -> tuple[str, str, int]:
    """Chunk + embed + upsert. Returns (namespace, doc_id, parent_count)."""
    parents = chunk_document(parsed)
    children = [c for p in parents for c in p.children]
    if not children:
        raise ValueError("No content could be extracted from the uploaded file.")

    embeddings = await dashscope_client.embed([c.text for c in children])
    collection = ensure_collection(settings.milvus_session_collection)
    namespace = f"{settings.session_doc_chunk_prefix}{thread_id}"
    fake_doc_id = f"sess_{thread_id}_{sha[:12]}"

    rows = [
        IndexableChunk(
            id=build_child_chunk_id(
                fake_doc_id, child.parent_index, child.child_index_in_parent
            ),
            text=child.text,
            dense=vec,
            source="session",
            doc_id=fake_doc_id,
            chunk_type="child",
            parent_index=child.parent_index,
            heading_path=child.heading_path[:500],
            era="",
            namespace=namespace,
        )
        for child, vec in zip(children, embeddings)
    ]
    upsert_chunks(collection, rows)
    return namespace, fake_doc_id, len(parents)


async def _optional_dashscope_upload(storage_path: Path) -> str | None:
    try:
        return await dashscope_client.upload_file(storage_path)
    except Exception as e:
        logger.warning("DashScope Files upload skipped/failed for {}: {}", storage_path.name, e)
        return None


async def process_session_file_row(session: AsyncSession, row: SessionFile) -> SessionFile:
    """Ingest an existing SessionFile row (status processing → ready/failed)."""
    if not row.storage_path:
        raise ValueError("Missing storage_path")
    storage_path = Path(row.storage_path)
    if not storage_path.is_file():
        raise ValueError(f"Storage file missing: {row.storage_path}")

    original_filename = row.filename
    sha = row.file_sha256 or _file_sha(storage_path)
    size = row.size_bytes or storage_path.stat().st_size
    ext = storage_path.suffix.lower()
    if ext not in SESSION_UPLOAD_EXTENSIONS:
        raise ValueError(f"Unsupported session upload type: {ext}")
    if ext in IMAGE_EXTENSIONS and size > settings.session_image_max_bytes:
        raise ValueError(
            f"Image too large ({size} bytes); max {settings.session_image_max_bytes}"
        )

    # Drop previous vectors on retry
    if row.milvus_namespace or (row.extra_metadata or {}).get("doc_id"):
        try:
            await purge_session_file_vectors(row)
        except Exception as e:
            logger.warning("Failed to purge old session vectors for {}: {}", row.id, e)

    parsed, parser_used = await _parse_session_document(
        storage_path, ext=ext, original_filename=original_filename
    )
    full = parsed.full_text().strip()
    if ext in ({".pdf", ".docx"} | IMAGE_EXTENSIONS) and len(full) < settings.session_min_extract_chars:
        raise ValueError(
            f"Extracted text too short ({len(full)} chars) for {original_filename}."
        )

    namespace, doc_id, parent_count = await _index_parsed_to_session(
        thread_id=row.thread_id, sha=sha, parsed=parsed
    )

    use_fulltext = _wants_fulltext(storage_path, full, ext)
    mode = "fulltext" if use_fulltext else "session_rag"

    max_chars = _inline_max_chars()
    meta: dict = {
        "doc_id": doc_id,
        "title": parsed.title,
        "parser": parser_used,
    }
    if use_fulltext:
        meta["extracted_text"] = full[:max_chars]
        meta["extracted_truncated"] = len(full) > max_chars

    dashscope_file_id = row.dashscope_file_id
    if use_fulltext:
        # Strategy A: keep Files API upload as backup; failure does not block ready.
        uploaded = await _optional_dashscope_upload(storage_path)
        if uploaded:
            dashscope_file_id = uploaded

    row.file_sha256 = sha
    row.size_bytes = size
    row.mime_type = ext.lstrip(".")
    row.mode = mode
    row.dashscope_file_id = dashscope_file_id
    row.milvus_namespace = namespace
    row.chunks_count = parent_count
    row.status = "ready"
    row.extra_metadata = meta
    await session.commit()
    await session.refresh(row)
    logger.info(
        "Session file {} ready mode={} parser={} chunks={}",
        original_filename,
        mode,
        parser_used,
        parent_count,
    )
    return row


async def ingest_session_file(
    session: AsyncSession,
    *,
    thread_id: str,
    storage_path: Path,
    original_filename: str,
    existing: SessionFile | None = None,
) -> SessionFile:
    """Create (or reuse) a row and process it. Prefer async upload path with existing row."""
    sha = _file_sha(storage_path)
    size = storage_path.stat().st_size
    ext = storage_path.suffix.lower()
    if existing is None:
        row = SessionFile(
            thread_id=thread_id,
            filename=original_filename,
            storage_path=str(storage_path),
            file_sha256=sha,
            size_bytes=size,
            mime_type=ext.lstrip("."),
            mode="pending",
            status="processing",
            chunks_count=0,
        )
        session.add(row)
        await session.commit()
        await session.refresh(row)
    else:
        row = existing
        row.storage_path = str(storage_path)
        row.filename = original_filename
        row.file_sha256 = sha
        row.size_bytes = size
        row.mime_type = ext.lstrip(".")
        row.status = "processing"
        row.mode = "pending"
        meta = dict(row.extra_metadata or {})
        meta.pop("error", None)
        row.extra_metadata = meta or None
        await session.commit()
        await session.refresh(row)

    return await process_session_file_row(session, row)


async def mark_session_file_failed(session: AsyncSession, row: SessionFile, error: str) -> SessionFile:
    meta = dict(row.extra_metadata or {})
    meta["error"] = error[:2000]
    row.extra_metadata = meta
    row.status = "failed"
    if row.mode in {"pending", ""}:
        row.mode = "pending"
    await session.commit()
    await session.refresh(row)
    return row


async def purge_session_file_vectors(row: SessionFile) -> None:
    from app.knowledge.indexer import drop_doc, drop_namespace

    collection = settings.milvus_session_collection
    if row.mode not in _INDEXED_MODES and not row.milvus_namespace:
        return
    meta = row.extra_metadata or {}
    doc_id = meta.get("doc_id")
    if doc_id:
        await asyncio.to_thread(drop_doc, collection, str(doc_id))
    elif row.milvus_namespace:
        await asyncio.to_thread(drop_namespace, collection, row.milvus_namespace)


async def purge_thread_session_resources(
    session: AsyncSession, thread_id: str, rows: Iterable[SessionFile] | None = None
) -> None:
    if rows is None:
        rows = (
            await session.execute(select(SessionFile).where(SessionFile.thread_id == thread_id))
        ).scalars().all()
    for row in rows:
        if row.dashscope_file_id:
            try:
                await dashscope_client.delete_file(row.dashscope_file_id)
            except Exception as e:
                logger.warning("Failed to delete DashScope file {}: {}", row.dashscope_file_id, e)
        if row.mode in _INDEXED_MODES or row.milvus_namespace:
            await purge_session_file_vectors(row)


async def build_session_system_messages(
    session: AsyncSession, thread_id: str
) -> list[dict[str, str]]:
    rows = (
        await session.execute(
            select(SessionFile).where(
                SessionFile.thread_id == thread_id,
                SessionFile.mode.in_(list(_FULLTEXT_MODES)),
                SessionFile.status == "ready",
            )
        )
    ).scalars().all()
    messages: list[dict[str, str]] = []
    dirty = False
    capable = _fileid_capable()
    max_chars = _inline_max_chars()
    for row in rows:
        if capable and row.dashscope_file_id:
            messages.append({"role": "system", "content": f"fileid://{row.dashscope_file_id}"})
        meta = dict(row.extra_metadata or {})
        if not (meta.get("extracted_text") or "").strip() and row.storage_path:
            text = _extract_attachment_text(Path(row.storage_path))
            if text:
                meta["extracted_text"] = text[:max_chars]
                meta["extracted_truncated"] = len(text) > max_chars
                row.extra_metadata = meta
                dirty = True
        inline = _inline_text_attachment(row)
        if inline:
            messages.append(inline)
    if dirty:
        await session.commit()
    return messages


def _inline_text_attachment(row: SessionFile) -> dict[str, str] | None:
    """Inject extracted body for fulltext/files_api attachments."""
    meta = row.extra_metadata if isinstance(row.extra_metadata, dict) else {}
    text = (meta.get("extracted_text") or "").strip() if meta else ""
    if not text and row.storage_path:
        text = (_extract_attachment_text(Path(row.storage_path)) or "").strip()
    if not text:
        return None
    max_chars = _inline_max_chars()
    clipped = text[:max_chars]
    suffix = "\n\n（附件内容已截断）" if len(text) > max_chars else ""
    return {
        "role": "system",
        "content": f"会话附件《{row.filename}》内容：\n{clipped}{suffix}",
    }


async def session_namespace_filter(
    session: AsyncSession, thread_id: str
) -> str | None:
    rows = (
        await session.execute(
            select(SessionFile).where(
                SessionFile.thread_id == thread_id,
                SessionFile.status == "ready",
                SessionFile.mode.in_(list(_INDEXED_MODES)),
            )
        )
    ).scalars().all()
    ns_values = [r.milvus_namespace for r in rows if r.milvus_namespace]
    if not ns_values:
        return None
    items = ", ".join(f'"{v}"' for v in ns_values)
    return f"namespace in [{items}]"


def get_preview_text(row: SessionFile) -> tuple[str, bool]:
    """Return (text, truncated) for preview API."""
    meta = row.extra_metadata if isinstance(row.extra_metadata, dict) else {}
    text = (meta.get("extracted_text") or "").strip() if meta else ""
    truncated = bool(meta.get("extracted_truncated")) if meta else False
    if text:
        return text, truncated
    if row.storage_path:
        extracted = _extract_attachment_text(Path(row.storage_path)) or ""
        max_chars = _inline_max_chars()
        if len(extracted) > max_chars:
            return extracted[:max_chars], True
        return extracted, False
    return "", False
